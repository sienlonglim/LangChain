import streamlit as st
import re
import yaml
from io import StringIO
from PyPDF2 import PdfReader
from docx import Document as docx
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
from langchain.embeddings.openai import OpenAIEmbeddings
from langchain.chat_models import ChatOpenAI
from langchain.prompts import PromptTemplate
from langchain.chains import RetrievalQA
from langchain.vectorstores import FAISS

def get_pdf(file : bytes, pdf_title : str, config : dict):
    use_splitter = config['splitter_options']['use_splitter']
    remove_leftover_delimiters = config['splitter_options']['remove_leftover_delimiters']
    remove_pages = config['splitter_options']['remove_pages']
    chunk_size = config['splitter_options']['chunk_size']
    chunk_overlap = config['splitter_options']['chunk_overlap']
    chunk_separators = config['splitter_options']['chunk_separators']
    front_pages_to_remove = config['splitter_options']['front_pages_to_remove']
    last_pages_to_remove = config['splitter_options']['last_pages_to_remove']
    delimiters_to_remove = config['splitter_options']['delimiters_to_remove']

    reader = PdfReader(file)
    for key, value in reader.metadata.items():
        if 'title' in key.lower():
            pdf_title = value
    print(f'\t\tOriginal no. of pages: {len(reader.pages)}')

    # Remove pages
    if remove_pages:
        for _ in range(front_pages_to_remove):
            del reader.pages[0]
        for _ in range(last_pages_to_remove):
            reader.pages.pop()
        print(f'\tNumber of pages after skipping: {len(reader.pages)}')
    
    # Each file will be split into LangChain Documents and kept in a list
    document_chunks = []

    # Loop through each page and extract the text and write in metadata
    if use_splitter:
        splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size,
                                            chunk_overlap=chunk_overlap,
                                            separators = chunk_separators)
        for page_number, page in enumerate(reader.pages):
            page_chunks = splitter.split_text(page.extract_text()) # splits into a list of strings (chunks)

            # Write into Document format with metadata
            for chunk in page_chunks:
                document_chunks.append(Document(page_content=chunk, metadata={'source': pdf_title , 'page':str(page_number+1)}))
        
    else:
        # This will split purely by page
        for page_number, page in enumerate(reader.pages):
            document_chunks.append(Document(page_content=page.extract_text(),  metadata={'source': pdf_title , 'page':str(page_number+1)}))
            i += 1

    # Remove all left over the delimiters and extra spaces
            if remove_leftover_delimiters:
                for chunk in document_chunks:
                    for delimiter in delimiters_to_remove:
                        chunk.page_content = re.sub(delimiter, ' ', chunk.page_content)

    print(f'\t\tExtracted no. of documents: {len(document_chunks)}')
    return pdf_title, document_chunks

def get_txt(file : bytes, title : str, config : dict):
    use_splitter = config['splitter_options']['use_splitter']
    remove_leftover_delimiters = config['splitter_options']['remove_leftover_delimiters']
    chunk_size = config['splitter_options']['chunk_size']
    chunk_overlap = config['splitter_options']['chunk_overlap']
    chunk_separators = config['splitter_options']['chunk_separators']
    delimiters_to_remove = config['splitter_options']['delimiters_to_remove']

    # file.getvalue() returns bytes which is decoded using utf-8, to instantiate a StringIO. Then we getvalue of the StringIO
    data = StringIO(file.getvalue().decode("utf-8"))
    text = data.getvalue()

    if use_splitter:
        splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size,
                                                chunk_overlap=chunk_overlap,
                                                separators = chunk_separators)
        splits = splitter.split_text(text) 
        document_chunks = []
        for split in splits:
            document_chunks.append((Document(page_content=split,  metadata={'source': title , 'page':'na'})))
    else:
        document_chunks = [(Document(page_content=text,  metadata={'source': title , 'page':'na'}))]
    
    # Remove all left over the delimiters and extra spaces
    if remove_leftover_delimiters:
        for chunk in document_chunks:
            for delimiter in delimiters_to_remove:
                chunk.page_content = re.sub(delimiter, ' ', chunk.page_content)

    print(f'\t\tExtracted no. of documents: {len(document_chunks)}')
    return title, document_chunks

def get_docx(file : bytes, title : str, config : dict):
    use_splitter = config['splitter_options']['use_splitter']
    remove_leftover_delimiters = config['splitter_options']['remove_leftover_delimiters']
    chunk_size = config['splitter_options']['chunk_size']
    chunk_overlap = config['splitter_options']['chunk_overlap']
    chunk_separators = config['splitter_options']['chunk_separators']
    delimiters_to_remove = config['splitter_options']['delimiters_to_remove']


    word_doc = docx(file)
    text = ''
    for paragraph in word_doc.paragraphs:
        text += paragraph.text

    if use_splitter:
        splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size,
                                                chunk_overlap=chunk_overlap,
                                                separators = chunk_separators)
        splits = splitter.split_text(text) 
        document_chunks = []
        for split in splits:
            document_chunks.append((Document(page_content=split,  metadata={'source': title , 'page':'na'})))
    else:
        document_chunks = [(Document(page_content=text,  metadata={'source': title , 'page':'na'}))]
    
    # Remove all left over the delimiters and extra spaces
    if remove_leftover_delimiters:
        for chunk in document_chunks:
            for delimiter in delimiters_to_remove:
                chunk.page_content = re.sub(delimiter, ' ', chunk.page_content)

    print(f'\t\tExtracted no. of documents: {len(document_chunks)}')
    return title, document_chunks

def get_chunks(file_input : list, config : dict):
    # Main list of all LangChain Documents
    document_chunks_full = []
    document_names = []
    print(f'Splitting documents: total of {len(file_input)}')

    # Handle file by file
    for file_index, file in enumerate(file_input):

        # Get the file type and file name
        file_type = file.type.split('/')[1]
        print(f'\tSplitting file {file_index+1} : {file.name}')
        file_name = file.name.split('.')[:-1]
        file_name = ''.join(file_name)

        # Handle different file types
        if file_type =='pdf':
            title, document_chunks = get_pdf(file, file_name, config)
        elif file_type == 'txt' or file_type == 'plain':
            title, document_chunks = get_txt(file, file_name, config)
        elif file_type == 'vnd.openxmlformats-officedocument.wordprocessingml.document':
            title, document_chunks = get_docx(file, file_name, config)
        
        document_names.append(title)
        document_chunks_full.extend(document_chunks)
             
    print(f'\tNumber of document chunks extracted in total: {len(document_chunks_full)}')
    return document_chunks_full, document_names

def get_embeddings(openai_api_key : str, documents : list,  config : dict):
    # create the open-source embedding function
    model = config['embedding_options']['model']
    db_option = config['embedding_options']['db_option']
    # persist_directory = config['embedding_options']['persist_directory']

    embedding_function = OpenAIEmbeddings(deployment="SL-document_embedder",
                                        model=model,
                                        show_progress_bar=True,
                                        openai_api_key = openai_api_key) 

    # Load it into FAISS
    print('Initializing vector_db')
    if db_option == 'FAISS':
        print('\tRunning in memory')
        vector_db = FAISS.from_documents(documents = documents, 
                                        embedding = embedding_function)
    print('\tCompleted')

    # If successful, increment the usage based on number of documents
    if openai_api_key == st.secrets['openai_api_key']:
        docs_processed = len(st.session_state.doc_names)
        increment_api_usage_count(docs_processed)
        print(f'Current usage_counter: {st.session_state.usage_counter}')
    return vector_db

def get_llm(openai_api_key : str, temperature : int, config : dict):
    model_name = config['llm']
    # Instantiate the llm object 
    print('Instantiating the llm')
    try:
        llm = ChatOpenAI(model_name=model_name, 
                        temperature=temperature, 
                        api_key=openai_api_key)
    except Exception as e:
        print(e)
        st.error('Error occured, check that your API key is correct.', icon="🚨")
    else:
        print('\tCompleted')
    return llm 

def get_chain(llm : object, vector_db : object, prompt_mode : str):
    print('Creating chain from template')
    if prompt_mode == 'Restricted':
        # Build prompt template
        template = """Use the following pieces of context to answer the question at the end. \
        If you don't know the answer, just say that you don't know, don't try to make up an answer. \
        Keep the answer as concise as possible. 
        Context: {context}
        Question: {question}
        Helpful Answer:"""
        qa_chain_prompt = PromptTemplate.from_template(template)
    elif prompt_mode == 'Creative':
        # Build prompt template
        template = """Use the following pieces of context to answer the question at the end. \
        If you don't know the answer, you may make inferences, but make it clear in your answer. \
        Keep the answer as concise as possible. 
        Context: {context}
        Question: {question}
        Helpful Answer:"""
        qa_chain_prompt = PromptTemplate.from_template(template)

    # Build QuestionAnswer chain
    qa_chain = RetrievalQA.from_chain_type(
        llm,
        retriever=vector_db.as_retriever(),
        return_source_documents=True,
        chain_type_kwargs={"prompt": qa_chain_prompt}
        )
    print('\tCompleted')
    return qa_chain

def get_response(user_input : str, qa_chain : object):
    # Query and Response
    print('Getting response from server')
    result = qa_chain({"query": user_input})
    return result

def get_settings():
    '''
    Handles initializing of session_state variables
    '''
    # Create empty list of document names
    if 'doc_names' not in st.session_state:
        st.session_state.doc_names = None
    
    # Load config if not yet loaded
    if 'config' not in st.session_state:
        with open('config.yml', 'r') as file:
            st.session_state.config = yaml.safe_load(file)

    # Create an API call counter, to cap usage
    if 'usage_counter' not in st.session_state:
        st.session_state.usage_counter = 0

    # Check if host api key is enabled
    if st.session_state.config['enable_host_api_key']:
        openai_api_key = st.secrets['openai_api_key']
    else:
        openai_api_key = 'NA'
    return openai_api_key

def increment_api_usage_count(increment : int):
    st.session_state.usage_counter += increment

def main():
    '''
    Main Function
    '''
    # Load configs and check for session_states
    st.set_page_config(page_title="Document Query Bot ")
    result = None
    openai_api_key_host = get_settings()

    # Sidebar
    with st.sidebar:
        # API option
        if st.session_state.usage_counter >= 5:
            api_option = st.radio('API key usage', ['Host API key usage cap reached!'])
        else:
            api_option = st.radio('API key usage', ['Use my own API key', 'Use host API key (capped)'], help='Cap is counted by number of documents uploaded (max 5 in total)')
        if (api_option == 'Use my own API key') or (api_option == 'Host API key usage cap reached!'):
            openai_api_key =st.text_input("Enter your API key")
        else:
            openai_api_key = openai_api_key_host
        
        # Document uploader
        documents = st.file_uploader(label = 'Upload documents for embedding to VectorDB', 
                                    help = 'Overwrites any existing files uploaded',
                                    type = ['pdf', 'txt', 'docx'], 
                                    accept_multiple_files=True)
        if st.button('Upload', type='primary') and documents:
            with st.status('Uploading... (this may take a while)', expanded=True) as status:
                try:
                    st.write("Splitting documents...")
                    document_chunks_full, st.session_state.doc_names = get_chunks(documents, st.session_state.config)
                    st.write("Creating embeddings...")
                    st.session_state.vector_db = get_embeddings(openai_api_key, document_chunks_full, st.session_state.config)
                except Exception as e:
                    print(e)
                    status.update(label='Error occured.', state='error', expanded=False)
                else:
                    status.update(label='Embedding complete!', state='complete', expanded=False)

    # Main page area
    st.markdown("### :rocket: Welcome to Sien Long's Document Query Bot")

    # Info bar
    if st.session_state.doc_names:
        doc_name_display = ''
        for doc_count, doc_name in enumerate(st.session_state.doc_names):
            doc_name_display += str(doc_count+1) + '. ' + doc_name + '\n\n'
    else:
        doc_name_display = 'No documents uploaded yet!'
    st.info(f"Current loaded document(s): \n\n {doc_name_display}", icon='ℹ️')
    if not openai_api_key:
        st.write('Enter your API key on the sidebar to begin')

    # Query form and response
    st.divider()
    with st.form('my_form'):
        user_input = st.text_area('Enter prompt:', value='Enter prompt here')

        # Select for model and prompt template settings
        if st.session_state.doc_names:
            prompt_mode = st.selectbox('Choose mode of prompt', ('Restricted', 'Creative'), 
                                    help='Restricted mode will reduce chances of LLM answering using out of context knowledge')
            temperature = st.select_slider('Select temperature', options=[x / 10 for x in range(0, 21)], 
                                        help='0 is recommended for restricted mode, 1 for creative mode. \n\
                                            Going above 1 creates more unpredictable responses and takes longer.')
        
        # Display error if no API key given
        if not openai_api_key.startswith('sk-'):
            if openai_api_key == 'NA':
                st.warning('Host key currently not available, please use your own OpenAI API key!', icon='⚠')
            else:
                st.warning('Please enter your OpenAI API key!', icon='⚠')

        # Submit a prompt
        if st.form_submit_button('Submit', type='primary') and openai_api_key.startswith('sk-'):
            with st.spinner('Loading...'):
                try:
                    st.session_state.llm = get_llm(openai_api_key, temperature, st.session_state.config)
                    st.session_state.qa_chain = get_chain(st.session_state.llm, st.session_state.vector_db, prompt_mode)
                    result = get_response(user_input, st.session_state.qa_chain)
                except Exception as e:
                    print(e)
                    st.error('Error occured, unable to process response!', icon="🚨")

            if result:
                # Display the result
                st.info('Query Response:', icon='📕')
                st.info(result["result"])
                st.write(' ')
                st.info('Sources', icon='📚')
                for document in result['source_documents']:
                    st.write(document.page_content + '\n\n' + document.metadata['source'] + ' (pg ' + document.metadata['page'] + ')')
                    st.write('-----------------------------------')
                print('\tCompleted')

# Main
if __name__ == '__main__':
   main()
